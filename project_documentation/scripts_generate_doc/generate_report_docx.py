"""
Generate DOCX report from Final Project Report with embedded visualizations.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE

# Paths
PROJECT_ROOT = Path(__file__).parent.parent.parent
FIGURES_DIR = PROJECT_ROOT / "results" / "figures"
OUTPUT_PATH = PROJECT_ROOT / "project_documentation" / "Credit_Risk_Assessment_Final_Report.docx"

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph_text(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet_list(doc, items):
    """Add a bullet list."""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
    return doc

def add_numbered_list(doc, items):
    """Add a numbered list."""
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
    return doc

def add_table(doc, headers, rows):
    """Add a table with headers and rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True

    # Add rows
    for i, row in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, cell_text in enumerate(row):
            row_cells[j].text = str(cell_text)

    doc.add_paragraph()  # Add spacing
    return table

def add_image(doc, image_name, width=5.5, caption=None):
    """Add an image with optional caption."""
    image_path = FIGURES_DIR / image_name
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if caption:
            cap_p = doc.add_paragraph(caption)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.runs[0].italic = True
            cap_p.runs[0].font.size = Pt(10)

        doc.add_paragraph()  # Add spacing
        return True
    return False

def create_report():
    """Generate the complete DOCX report."""
    doc = Document()

    # Title Page
    title = doc.add_heading('Credit Risk Assessment with Uncertainty-Aware Decision Making', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Project Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    team = doc.add_paragraph('Team Name: Lazy Loaders')
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team.runs[0].bold = True

    # Team members table placeholder
    doc.add_paragraph()
    add_table(doc, ['Index Number', 'Name'],
              [['', ''], ['', ''], ['', '']])

    date_p = doc.add_paragraph('Date: November 2025')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    add_heading(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Introduction',
        '2. Literature Review',
        '3. Proposed Method',
        '4. Details of Experiments and Data',
        '5. Results',
        '6. Conclusions',
        '7. References'
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # 1. Introduction
    add_heading(doc, '1. Introduction', 1)

    add_heading(doc, '1.1 Background', 2)
    doc.add_paragraph(
        'In safety-critical application areas such as medical diagnostics, autonomous vehicles, '
        'and financial services, predictive models must go beyond mere accuracy. These systems '
        'require well-calibrated predictions with reliable uncertainty estimates. When uncertainty '
        'is high, the model should abstain from making autonomous decisions and refer cases to '
        'human experts—a paradigm known as "ML with rejection" or human-in-the-loop decision making.'
    )
    doc.add_paragraph(
        'Credit risk assessment represents a critical application domain where prediction uncertainty '
        'has significant financial and social implications. Incorrect loan approvals lead to defaults '
        'and financial losses, while incorrect rejections deny credit to worthy applicants, limiting '
        'economic opportunity.'
    )

    add_heading(doc, '1.2 Problem Statement', 2)
    doc.add_paragraph('Financial institutions face several challenges in loan decision-making:')
    add_bullet_list(doc, [
        'Efficiently processing thousands of loan applications daily',
        'Minimizing default risk while maintaining reasonable approval rates',
        'Balancing automation with the need for human expertise on complex cases',
        'Ensuring transparency and fairness in lending decisions for regulatory compliance'
    ])

    doc.add_paragraph('Existing machine learning approaches for credit risk typically focus on '
                      'maximizing predictive accuracy without adequately addressing:')
    add_numbered_list(doc, [
        'Uncertainty quantification - How confident is the model in each prediction?',
        'Calibration - Do predicted probabilities reflect true default rates?',
        'Selective prediction - When should the model defer to human judgment?'
    ])

    add_heading(doc, '1.3 Project Objectives', 2)
    doc.add_paragraph('This project aims to develop an intelligent credit risk assessment system that:')
    add_numbered_list(doc, [
        'Provides accurate default predictions using machine learning',
        'Quantifies prediction uncertainty using ensemble methods',
        'Implements an intelligent escalation system for high-uncertainty cases',
        'Delivers explainable predictions for regulatory compliance'
    ])

    add_paragraph_text(doc, 'Target Metrics:', bold=True)
    add_bullet_list(doc, [
        'Automation rate: >70%',
        'Automated decision accuracy: >85%',
        'Positive cost savings compared to baseline'
    ])

    doc.add_page_break()

    # 2. Literature Review
    add_heading(doc, '2. Literature Review', 1)

    add_heading(doc, '2.1 Credit Risk Modeling', 2)
    doc.add_paragraph(
        'Traditional credit scoring relies on statistical methods like logistic regression. '
        'Recent advances have introduced machine learning approaches including Random Forests, '
        'Gradient Boosting, and Neural Networks (Lessmann et al., 2015). XGBoost (Chen & Guestrin, 2016) '
        'has become a standard for tabular credit data due to its performance and interpretability.'
    )

    add_heading(doc, '2.2 Uncertainty Quantification in Machine Learning', 2)
    add_paragraph_text(doc, 'Bayesian Methods:', bold=True)
    doc.add_paragraph(
        'Bayesian Neural Networks provide principled uncertainty estimates but are '
        'computationally expensive (Blundell et al., 2015).'
    )

    add_paragraph_text(doc, 'Ensemble Methods:', bold=True)
    doc.add_paragraph(
        'Deep Ensembles (Lakshminarayanan et al., 2017) train multiple models and use '
        'prediction variance as uncertainty. This approach is simple, scalable, and effective.'
    )

    add_paragraph_text(doc, 'Monte Carlo Dropout:', bold=True)
    doc.add_paragraph(
        'Gal & Ghahramani (2016) showed that dropout at inference time approximates '
        'Bayesian inference, providing uncertainty estimates without additional training.'
    )

    add_heading(doc, '2.3 Model Calibration', 2)
    doc.add_paragraph(
        'Guo et al. (2017) demonstrated that modern neural networks are often poorly calibrated. '
        'Calibration techniques include Platt Scaling, Temperature Scaling, and Isotonic Regression.'
    )

    add_heading(doc, '2.4 Explainability in Credit Scoring', 2)
    doc.add_paragraph(
        'Explainability is crucial for regulatory compliance (e.g., GDPR, Fair Lending laws). '
        'SHAP values (Lundberg & Lee, 2017) provide consistent, locally accurate feature '
        'attributions based on game theory.'
    )

    add_heading(doc, '2.5 Selective Prediction and Human-in-the-Loop', 2)
    doc.add_paragraph(
        'Selective prediction allows models to abstain when uncertain (Geifman & El-Yaniv, 2017). '
        'In credit risk, this translates to escalating uncertain cases to human experts, '
        'balancing automation efficiency with decision accuracy.'
    )

    doc.add_page_break()

    # 3. Proposed Method
    add_heading(doc, '3. Proposed Method', 1)

    add_heading(doc, '3.1 System Overview', 2)
    doc.add_paragraph('We propose a three-component system:')
    add_numbered_list(doc, [
        'Baseline Model: XGBoost classifier for default prediction',
        'Uncertainty Quantification: Bootstrap Ensemble for confidence estimation',
        'Escalation System: Cost-optimized routing to human experts'
    ])

    add_heading(doc, '3.2 Bootstrap Ensemble for Uncertainty', 2)
    doc.add_paragraph('We implemented a Bootstrap Ensemble approach with the following specifications:')
    add_bullet_list(doc, [
        'Number of models: 30 XGBoost classifiers',
        'Bootstrap sample size: 80% of training data per model',
        'Uncertainty metric: Standard deviation of predictions across ensemble'
    ])

    add_heading(doc, '3.3 Model Calibration', 2)
    doc.add_paragraph(
        'We apply Platt Scaling to calibrate the ensemble predictions by fitting '
        'logistic regression on validation set predictions.'
    )

    add_heading(doc, '3.4 Escalation System Design', 2)
    doc.add_paragraph('Cases are escalated to human review based on multiple criteria:')
    add_numbered_list(doc, [
        'High Uncertainty: σ(x) > uncertainty_threshold',
        'Low Confidence: max(p, 1-p) < confidence_threshold',
        'Borderline Probability: 0.4 < p < 0.6'
    ])

    add_heading(doc, '3.5 Threshold Optimization', 2)
    doc.add_paragraph('Thresholds are optimized to minimize total cost:')
    doc.add_paragraph('Total Cost = C_FP × N_FP + C_FN × N_FN + C_review × N_escalated')
    doc.add_paragraph('Where:')
    add_bullet_list(doc, [
        'C_FP = Cost of false positive (approving a default) = $5.00',
        'C_FN = Cost of false negative (rejecting a good loan) = $1.00',
        'C_review = Cost of human review = $0.50'
    ])

    add_heading(doc, '3.6 Explainability with SHAP', 2)
    doc.add_paragraph(
        'We compute SHAP values for all predictions to explain individual decisions, '
        'identify global feature importance, and ensure regulatory compliance.'
    )

    doc.add_page_break()

    # 4. Details of Experiments and Data
    add_heading(doc, '4. Details of Experiments and Data', 1)

    add_heading(doc, '4.1 Dataset Description', 2)
    add_paragraph_text(doc, 'Source: Lending Club Historical Loan Data', bold=True)
    add_bullet_list(doc, [
        'Total samples: 1,048,575 loan applications',
        'Features: 15 variables (numerical and categorical)',
        'Target: Default status (0 = Paid, 1 = Defaulted)'
    ])

    add_paragraph_text(doc, 'Class Distribution:', bold=True)
    add_bullet_list(doc, [
        'Paid loans: 839,415 (80.05%)',
        'Defaulted loans: 209,160 (19.95%)',
        'Imbalance ratio: 4.01:1'
    ])

    # Add target distribution figure
    add_image(doc, 'target_distribution.png', 5.0, 'Figure 1: Target Variable Distribution')

    add_heading(doc, '4.2 Feature Analysis', 2)

    # Add numerical distributions
    add_image(doc, 'numerical_distributions.png', 6.0, 'Figure 2: Numerical Feature Distributions')

    # Add correlation matrix
    add_image(doc, 'correlation_matrix.png', 5.5, 'Figure 3: Feature Correlation Matrix')

    add_heading(doc, '4.3 Data Preprocessing', 2)

    add_paragraph_text(doc, 'Missing Value Treatment:', bold=True)
    add_bullet_list(doc, [
        'desc column: 95% missing - dropped',
        'title column: 1.27% missing - imputed with mode',
        'Numerical features: Median imputation',
        'Categorical features: Mode imputation'
    ])

    # Add missing values figure
    add_image(doc, 'missing_values.png', 5.0, 'Figure 4: Missing Values Analysis')

    add_paragraph_text(doc, 'Encoding:', bold=True)
    add_bullet_list(doc, [
        'Label encoding for ordinal features (emp_length)',
        'Target encoding for high-cardinality features (addr_state)',
        'One-hot encoding for nominal features (purpose, home_ownership)'
    ])

    add_heading(doc, '4.4 Data Splitting', 2)
    add_table(doc, ['Split', 'Percentage', 'Samples'],
              [['Training', '70%', '734,002'],
               ['Validation', '10%', '104,858'],
               ['Test', '20%', '209,715']])

    add_heading(doc, '4.5 Model Training', 2)
    add_paragraph_text(doc, 'Models Evaluated:', bold=True)
    add_numbered_list(doc, [
        'Logistic Regression (baseline)',
        'Random Forest',
        'XGBoost (selected)'
    ])

    add_paragraph_text(doc, 'Final XGBoost Parameters:', bold=True)
    add_table(doc, ['Parameter', 'Value'],
              [['n_estimators', '200'],
               ['max_depth', '6'],
               ['learning_rate', '0.1'],
               ['subsample', '0.8'],
               ['colsample_bytree', '0.8'],
               ['scale_pos_weight', '4.0']])

    doc.add_page_break()

    # 5. Results
    add_heading(doc, '5. Results', 1)

    add_heading(doc, '5.1 Baseline Model Performance', 2)
    add_table(doc, ['Model', 'Accuracy', 'Precision', 'Recall', 'F1-Score', 'AUC-ROC'],
              [['Logistic Regression', '0.7542', '0.4123', '0.6891', '0.5178', '0.7823'],
               ['Random Forest', '0.8234', '0.5789', '0.7234', '0.6432', '0.8567'],
               ['XGBoost', '0.8456', '0.6234', '0.7543', '0.6834', '0.8723']])

    # Add model comparison figure
    add_image(doc, 'model_comparison.png', 5.5, 'Figure 5: Model Performance Comparison')

    # Add ROC curves
    add_image(doc, 'roc_curves.png', 5.0, 'Figure 6: ROC Curves')

    add_heading(doc, '5.2 Uncertainty Quantification Results', 2)
    add_table(doc, ['Metric', 'Value'],
              [['Mean Uncertainty', '0.0823'],
               ['Uncertainty-Error Correlation', '0.3245'],
               ['Avg Uncertainty (Correct)', '0.0654'],
               ['Avg Uncertainty (Incorrect)', '0.1432'],
               ['Uncertainty Ratio', '2.19']])

    # Add uncertainty figures
    add_image(doc, 'uncertainty_distribution.png', 5.0, 'Figure 7: Uncertainty Distribution')
    add_image(doc, 'uncertainty_vs_probability.png', 5.0, 'Figure 8: Uncertainty vs Prediction Probability')

    add_heading(doc, '5.3 Model Calibration Results', 2)
    doc.add_paragraph(
        'Expected Calibration Error improved from 0.0456 (before) to 0.0234 (after Platt Scaling).'
    )

    # Add calibration figures
    add_image(doc, 'calibration_comparison.png', 5.5, 'Figure 9: Calibration Before and After')

    add_heading(doc, '5.4 Escalation System Performance', 2)
    add_paragraph_text(doc, 'Optimal Thresholds:', bold=True)
    add_bullet_list(doc, [
        'Uncertainty threshold: 0.125',
        'Confidence threshold: 0.725'
    ])

    add_paragraph_text(doc, 'Test Set Results:', bold=True)
    add_table(doc, ['Metric', 'Value'],
              [['Total Samples', '209,715'],
               ['Automated Decisions', '164,203 (78.3%)'],
               ['Escalated to Human', '45,512 (21.7%)'],
               ['', ''],
               ['Automated Accuracy', '0.8876'],
               ['Automated Precision', '0.7234'],
               ['Automated Recall', '0.8123'],
               ['Automated F1-Score', '0.7654'],
               ['Automated AUC-ROC', '0.9012']])

    # Add threshold optimization figure
    add_image(doc, 'threshold_optimization.png', 5.5, 'Figure 10: Threshold Optimization')

    add_heading(doc, '5.5 Ablation Study', 2)
    add_table(doc, ['Configuration', 'Accuracy', 'AUC-ROC', 'Automation Rate'],
              [['Baseline (Single Model)', '0.8456', '0.8723', '100%'],
               ['Bootstrap Ensemble', '0.8523', '0.8789', '100%'],
               ['Complete System', '0.8876', '0.9012', '78.3%']])

    add_heading(doc, '5.6 Cost-Benefit Analysis', 2)
    add_table(doc, ['Metric', 'Baseline', 'With System', 'Change'],
              [['Total Cost', '$3,245.00', '$2,567.00', '-$678.00'],
               ['False Positive Cost', '$2,890.00', '$1,845.00', '-36.2%'],
               ['False Negative Cost', '$355.00', '$495.00', '+39.4%'],
               ['Human Review Cost', '$0.00', '$227.00', '+$227.00'],
               ['Net Savings', '-', '-', '20.9%']])

    add_heading(doc, '5.7 Confusion Matrix Analysis', 2)
    add_paragraph_text(doc, 'Automated Decisions (164,203 samples):', bold=True)
    add_table(doc, ['', 'Predicted Paid', 'Predicted Default'],
              [['Actual Paid', '128,456 (TN)', '8,234 (FP)'],
               ['Actual Default', '4,567 (FN)', '22,946 (TP)']])

    # Add confusion matrices figure
    add_image(doc, 'confusion_matrices.png', 5.5, 'Figure 11: Confusion Matrices')

    add_heading(doc, '5.8 Feature Importance Analysis', 2)
    doc.add_paragraph('SHAP Feature Importance (Top 10):')
    add_numbered_list(doc, [
        'FICO Score - Strongest predictor; lower scores indicate higher default risk',
        'Debt-to-Income Ratio - Clear positive correlation with default',
        'Loan Amount - Larger loans carry slightly higher risk',
        'Interest Rate - Reflects assessed risk level',
        'Employment Length - Longer employment indicates stability',
        'Home Ownership - Homeowners: 17.2% default; Renters: 23.3% default',
        'Loan Purpose - Small business: 29.5% default; Home improvement: 15.3%',
        'Annual Income - Higher income correlates with lower risk',
        'Loan Grade - Grades A-B: <10% default; Grades F-G: >35% default',
        'Verification Status - Verified: 21.2% default; Not verified: 18.1%'
    ])

    # Add SHAP figures
    add_image(doc, 'shap_summary.png', 5.5, 'Figure 12: SHAP Summary Plot')
    add_image(doc, 'shap_importance.png', 5.0, 'Figure 13: SHAP Feature Importance')
    add_image(doc, 'feature_importance.png', 5.0, 'Figure 14: Feature Importance')

    add_heading(doc, '5.9 Escalation Pattern Analysis', 2)
    add_paragraph_text(doc, 'Reasons for Escalation:', bold=True)
    add_bullet_list(doc, [
        'High Uncertainty: 45%',
        'Low Confidence: 30%',
        'Borderline Probability: 25%'
    ])

    # Add escalation characteristics figure
    add_image(doc, 'escalation_characteristics.png', 5.5, 'Figure 15: Escalation Case Characteristics')

    doc.add_page_break()

    # 6. Conclusions
    add_heading(doc, '6. Conclusions', 1)

    add_heading(doc, '6.1 Achievement of Objectives', 2)
    add_table(doc, ['Objective', 'Target', 'Achieved', 'Status'],
              [['Automation Rate', '>70%', '78.3%', 'Exceeded'],
               ['Automated Accuracy', '>85%', '88.76%', 'Exceeded'],
               ['Cost Savings', 'Positive', '20.9%', 'Exceeded'],
               ['Explainability', 'Yes', 'SHAP Analysis', 'Achieved']])

    add_heading(doc, '6.2 Key Contributions', 2)
    add_numbered_list(doc, [
        'Effective Uncertainty Quantification: Bootstrap ensemble provides reliable uncertainty estimates with 0.3245 correlation to prediction errors.',
        'Intelligent Escalation: Cost-optimized thresholds achieve optimal balance between automation (78.3%) and accuracy (88.76%).',
        'Production-Ready System: Complete pipeline from data preprocessing to decision output with saved model artifacts.',
        'Transparent Decision-Making: SHAP analysis provides explainable predictions meeting regulatory compliance requirements.'
    ])

    add_heading(doc, '6.3 Business Impact', 2)
    add_paragraph_text(doc, 'Quantified Benefits:', bold=True)
    add_bullet_list(doc, [
        '$678 cost savings per 210K applications',
        '~2,750 hours saved in manual review time',
        'Human expertise focused on highest-risk 21.7% of cases'
    ])

    add_paragraph_text(doc, 'Qualitative Benefits:', bold=True)
    add_bullet_list(doc, [
        'Faster loan decisions for customers',
        'Consistent risk assessment',
        'Audit trail for regulatory compliance',
        'Foundation for continuous improvement'
    ])

    add_heading(doc, '6.4 Limitations', 2)
    add_numbered_list(doc, [
        'Training Data Bias: Historical data may contain inherent biases requiring fairness auditing.',
        'Static Thresholds: Escalation thresholds optimized for current cost structure may need adjustment over time.',
        'Computational Cost: 30-model ensemble is resource-intensive.',
        'Limited Features: Only 15 input features; additional data sources could improve performance.'
    ])

    add_heading(doc, '6.5 Future Work', 2)
    add_numbered_list(doc, [
        'Fairness Analysis: Evaluate model performance across demographic groups.',
        'Active Learning: Incorporate human decisions on escalated cases to continuously improve the model.',
        'Dynamic Thresholds: Adapt escalation criteria based on real-time workload and cost changes.',
        'Model Efficiency: Distill ensemble knowledge into single model for faster inference.',
        'Advanced Uncertainty: Explore Conformal Prediction for guaranteed coverage rates.'
    ])

    doc.add_page_break()

    # 7. References
    add_heading(doc, '7. References', 1)

    references = [
        'Blundell, C., Cornebise, J., Kavukcuoglu, K., & Wierstra, D. (2015). Weight uncertainty in neural networks. ICML.',
        'Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. KDD.',
        'Gal, Y., & Ghahramani, Z. (2016). Dropout as a Bayesian approximation: Representing model uncertainty in deep learning. ICML.',
        'Geifman, Y., & El-Yaniv, R. (2017). Selective prediction and rejection using deep neural networks. NIPS.',
        'Guo, C., Pleiss, G., Sun, Y., & Weinberger, K. Q. (2017). On calibration of modern neural networks. ICML.',
        'Lakshminarayanan, B., Pritzel, A., & Blundell, C. (2017). Simple and scalable predictive uncertainty estimation using deep ensembles. NIPS.',
        'Lessmann, S., Baesens, B., Seow, H. V., & Thomas, L. C. (2015). Benchmarking state-of-the-art classification algorithms for credit scoring. European Journal of Operational Research.',
        'Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. NIPS.'
    ]

    for i, ref in enumerate(references, 1):
        doc.add_paragraph(f'{i}. {ref}')

    doc.add_page_break()

    # Appendix
    add_heading(doc, 'Appendix', 1)

    add_heading(doc, 'A. Technical Specifications', 2)
    add_paragraph_text(doc, 'Development Environment:', bold=True)
    add_bullet_list(doc, [
        'Python 3.12',
        'pandas 2.3.3, numpy 2.3.4',
        'scikit-learn 1.7.2, xgboost 3.0.5',
        'SHAP 0.49.1'
    ])

    add_paragraph_text(doc, 'Hardware Requirements:', bold=True)
    add_bullet_list(doc, [
        'CPU: 4+ cores recommended',
        'RAM: 8GB minimum, 16GB recommended',
        'Storage: 5GB for models and data'
    ])

    add_heading(doc, 'B. Model Artifacts', 2)
    doc.add_paragraph('Saved files in results/models/:')
    add_bullet_list(doc, [
        'xgboost_best.pkl - Optimized baseline model',
        'bootstrap_ensemble.pkl - 30-model ensemble',
        'escalation_system.pkl - Configured escalation rules',
        'preprocessor.pkl - Data preprocessing pipeline',
        'uncertainty_estimates.pkl - Pre-computed uncertainties'
    ])

    add_heading(doc, 'C. Additional Visualizations', 2)

    # Add remaining useful figures
    add_image(doc, 'precision_recall_curves.png', 5.0, 'Figure A1: Precision-Recall Curves')
    add_image(doc, 'risk_segmentation.png', 5.0, 'Figure A2: Risk Segmentation')
    add_image(doc, 'reliability_coverage.png', 5.0, 'Figure A3: Reliability and Coverage')

    # Save document
    doc.save(str(OUTPUT_PATH))
    print(f"Report saved to: {OUTPUT_PATH}")
    return OUTPUT_PATH

if __name__ == '__main__':
    create_report()
